from docx import Document  # Used to create and manipulate Word documents
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Provides access to paragraph alignment options
from docx.shared import Inches  # Allows setting sizes, like image dimensions, in inches
from docx.oxml import OxmlElement  # Provides access to XML elements in docx files
from docx.oxml.ns import qn  # Used for handling XML namespaces in docx
import google.generativeai as genai  # Provides access to Google Generative AI services
import os  
from database import Recommendation_English, Recommendation_Arabic  # Used to fetch recommendations in English and Arabic
import re  # Used for working with regular expressions

# Configure the generative AI model using the API key.
genai.configure(api_key=os.environ['API_KEY'])

try:
    # Try initializing the generative model (Gemini 1.5 Pro version).
    model = genai.GenerativeModel('gemini-1.5-pro')
except Exception as e:
    # If there's an error during model initialization, print the error and set the model to None.
    print(f"Error initializing model: {e}")
    model = None

# Function to generate energy efficiency recommendations in English.
def generate_recommendations_english(data):
    # Check if the model was successfully initialized.
    if model is None:
        # Return an error message if the model initialization failed.
        return "Unable to generate recommendations due to model initialization error."

    # Create a prompt using the input data to request specific energy-saving recommendations.
    prompt = f"""
    Based on the following energy audit data, provide 5-7 specific recommendations for saving power and improving energy efficiency:

        Accommodation: {data['type_of_accommodation']}
        number of residents: {data['number_of_residents']}
        Year of Construction: {data['year_of_construction']}
        Number of Bedrooms: {data['number_of_bedrooms']}
        Number of Floors: {data['number_of_floors']}
        Outdoor Garden: {data['outdoor_garden']}
        Swimming Pool: {data['swimming_pool']}
        Air Conditioning Systems: {data['ac_systems']}
        Lighting: {data['lighting']}
        Water Taps: {data['water_taps']}
        Water Heaters: {data['water_heaters']}
        Other Notes: {data['other']}'

    Please provide actionable and specific recommendations, benefits and implementation. Format the recommendations as a bullet point list in the following format.
    AC-System:

    Lighting:

    Water Taps:

    Water Heaters:

    Other Observations:

    Do not write anything before this and don't bold any sentences/words with **. Do not use '-' to bullet instead use numbers and use alphabets for sub points and don’t use brackets anywhere
    """

    # Generate content using the configured generative model.
    response = model.generate_content(prompt)
    
    # Clean up the generated text by removing asterisks (*) and any text in brackets.
    cleaned_response = response.text.replace("*", "")
    cleaned_response = re.sub(r'\[.*?\]', '', cleaned_response) 

    # Return the cleaned response with actionable recommendations.
    return cleaned_response


# Function to set background shading for a table cell in a Word document.
# Arguments:
# cell  - The table cell object (docx.table._Cell) where shading is to be applied.
# color - The fill color for the cell's shading, specified as a hex code (e.g., 'FF0000' for red).

def set_cell_shading(cell, color):
    # Create a new XML element for the shading (background color).
    shading = OxmlElement('w:shd')
    
    # Set the 'w:fill' attribute of the shading element to the specified color.
    shading.set(qn('w:fill'), color)
    
    # Access the cell's XML properties and append the shading element to apply the background color.
    cell._element.get_or_add_tcPr().append(shading)


# Function to set the text direction of a paragraph to Left-to-Right (LTR) in a Word document.
# Argument:
# paragraph - The paragraph object (docx.text.paragraph.Paragraph) where the LTR direction will be applied.

def set_ltr(paragraph):
    # Access the XML element of the paragraph.
    p = paragraph._element
    
    # Get or add the paragraph properties (pPr) element to modify paragraph attributes.
    pPr = p.get_or_add_pPr()
    
    # Create a new XML element for setting the bidirectional (bidi) attribute.
    # The 'w:bidi' element controls text direction, with '0' indicating Left-to-Right (LTR) direction.
    bidi = OxmlElement('w:bidi')
    
    # Set the 'w:val' attribute of the bidi element to '0', which means LTR.
    bidi.set(qn('w:val'), '0')
    
    # Append the bidi element to the paragraph properties to apply the LTR setting.
    pPr.append(bidi)
    
    # Set the paragraph alignment explicitly to left (LTR alignment).
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


# Function to set the text direction of a paragraph to Right-to-Left (RTL).
# Argument:
# paragraph - The paragraph object where RTL will be applied.
def set_rtl(paragraph):
    # Access the XML element of the paragraph.
    p = paragraph._element
    
    # Get or add the paragraph properties (pPr) element to modify paragraph attributes.
    pPr = p.get_or_add_pPr()
    
    # Create a new XML element for the bidirectional (bidi) attribute, which controls text direction.
    bidi = OxmlElement('w:bidi')
    
    # Set the 'w:val' attribute of the bidi element to '1', which means Right-to-Left (RTL).
    bidi.set(qn('w:val'), '1')
    
    # Append the bidi element to the paragraph properties to apply the RTL setting.
    pPr.append(bidi)
    
    # Align the paragraph text to the left explicitly (even though it's RTL direction).
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Function to set borders around a table cell.
# Argument:
# cell - The table cell object where borders will be applied.
def set_borders(cell):
    # Access the XML element of the table cell.
    tc = cell._element
    
    # Get or add the table cell properties (tcPr) element to modify cell attributes.
    tcPr = tc.get_or_add_tcPr()
    
    # Create a new XML element for the cell borders.
    tcBorders = OxmlElement('w:tcBorders')
    
    # Loop through each side of the cell (top, left, bottom, right) to set the border.
    for border_name in ['top', 'left', 'bottom', 'right']:
        # Create a border element for the current side.
        border = OxmlElement(f'w:{border_name}')
        
        # Set the border style to 'single' (solid line).
        border.set(qn('w:val'), 'single')
        
        # Set the border thickness (size) to 4, representing a width of 1/2 point.
        border.set(qn('w:sz'), '4')
        
        # Set the border spacing to 0 (no extra space between text and border).
        border.set(qn('w:space'), '0')
        
        # Set the border color to black ('000000' hex code).
        border.set(qn('w:color'), '000000')
        
        # Append the border element to the cell borders.
        tcBorders.append(border)
    
    # Append the borders to the table cell properties (tcPr).
    tcPr.append(tcBorders)

# Function to set line spacing in a paragraph.
# Argument:
# paragraph - The paragraph object where the spacing will be applied.
def set_paragraph_spacing(paragraph):
    # Get or add the paragraph properties (pPr) element to modify spacing attributes.
    pPr = paragraph._element.get_or_add_pPr()
    
    # Create a new XML element for the paragraph's line spacing.
    spacing = OxmlElement('w:spacing')
    
    # Set the line spacing to 360 TWIPS, equivalent to 1.5 line spacing (240 TWIPS is single spacing).
    spacing.set(qn('w:line'), '360')  # 1.5 * 240 TWIPS
    
    # Append the spacing element to the paragraph properties.
    pPr.append(spacing)

# Function to add a hyperlink to a paragraph.
# Arguments:
# paragraph - The paragraph object where the hyperlink will be added.
# url       - The URL that the hyperlink points to.
# text      - The display text for the hyperlink.
def add_hyperlink(paragraph, url, text):
    # Create the hyperlink relation for the paragraph's document part.
    part = paragraph.part
    
    # Relate the URL to the paragraph, creating an external hyperlink relation.
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create a new XML element for the hyperlink.
    hyperlink = OxmlElement('w:hyperlink')
    
    # Set the relation ID (r:id) to associate the hyperlink with the external URL.
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run (text segment) for the hyperlink text.
    new_run = OxmlElement('w:r')
    
    # Create a new run properties (rPr) element to style the hyperlink text.
    rPr = OxmlElement('w:rPr')

    # Add a color element to the run properties for blue-colored text (standard for hyperlinks).
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Hex value for blue color
    
    # Append the color element to the run properties.
    rPr.append(color)

    # Add an underline element to the run properties for underlining the hyperlink text.
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')  # Single underline style
    
    # Append the underline element to the run properties.
    rPr.append(underline)

    # Append the run properties to the run (applies styling to the text).
    new_run.append(rPr)

    # Create a text element (w:t) to hold the display text for the hyperlink.
    text_element = OxmlElement('w:t')
    
    # Set the text element's content to the display text.
    text_element.text = text
    
    # Append the text element to the run (displays the hyperlink text).
    new_run.append(text_element)
    
    # Append the run (with text and styling) to the hyperlink element.
    hyperlink.append(new_run)

    # Append the hyperlink element to the paragraph's XML structure.
    paragraph._element.append(hyperlink)
    
    # Return the updated paragraph containing the hyperlink.
    return paragraph


def create_report_english(data, recommendations):
         # Create a new Word document
    doc = Document()

    # -------------------- Header --------------------

    # Add logos to the header
    section = doc.sections[0] # Get the first section of the document
    header = section.header # Get the header of the section

    logo_paragraph = header.paragraphs[0] # Get the first paragraph of the header
    logo_run = logo_paragraph.add_run() # Add a run to the paragraph (a run is a sequence of text with the same formatting)
    logo_run.add_picture(r"./rak.png", width=Inches(1.6)) # Add the first logo (RAK)
    logo_run.add_text(" " * 70)  # Adjust the number of spaces to control the gap between logos
    logo_run.add_picture(r"./mun.png", width=Inches(2.0)) # Add the second logo (Municipality)
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Align the logo paragraph to the center

    # -------------------- Title --------------------

    # Title
    title = doc.add_heading(('Manzili Energy Audit Service Report'), level=1) # Add the report title as a heading level 1
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Center the title

    # -------------------- Report Number --------------------

    # Report Number
    report_number_paragraph = doc.add_paragraph(f"{('Report Number')}: {data['report_number']}") # Add the report number
    report_number_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Center the report number

    # -------------------- Overview --------------------

    # Overview
    overview_heading = doc.add_heading(('Overview'), level=2) # Add the "Overview" heading
    overview_paragraph = doc.add_paragraph((
        "This report summarizes the results and recommendations following the energy audit conducted in your home as part of the Manzili home energy consultancy service in Ras Al Khaimah. The goal of the audit is to help reduce your electricity and water bills and make your home more comfortable and modern."
    )) # Add the overview text

    # -------------------- Audit Details --------------------

    # Audit Details
    audit_details_heading = doc.add_heading(('Audit Details'), level=2) # Add the "Audit Details" heading
    audit_details_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT # Align the heading to the left
    audit_table = doc.add_table(rows=1, cols=4) # Create a table with 1 row and 4 columns for audit details
    hdr_cells = audit_table.rows[0].cells # Get the header cells of the table
    hdr_cells[0].text = ('Item') # Set the text for the first header cell
    hdr_cells[1].text = ('Details') # Set the text for the second header cell
    hdr_cells[2].text = ('Item') # Set the text for the third header cell
    hdr_cells[3].text = ('Details') # Set the text for the fourth header cell

    # Format the header cells
    for cell in hdr_cells:
        set_cell_shading(cell, "D3D3D3")  # Set header cell color to grey
        for paragraph in cell.paragraphs:
            set_ltr(paragraph) # Set text direction to left-to-right (assuming helper function)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Center align the text
        set_borders(cell) # Set cell borders (assuming helper function)

    # Define the audit fields to be included in the table
    audit_fields = [
        ('date_of_audit', 'report_number'),
        ('homeowner', 'contact_number'),
        ('location', 'type_of_accommodation'),
        ('house_number', 'year_of_construction'),
        ('number_of_bedrooms', 'number_of_floors')
    ]

    # Populate the audit details table
    for field_pair in audit_fields:
        row_cells = audit_table.add_row().cells # Add a new row to the table
        row_cells[0].text = (field_pair[0].replace('_', ' ').title()) # Set the text for the first cell in the row
        row_cells[1].text = data[field_pair[0]] # Set the text for the second cell in the row
        row_cells[2].text = (field_pair[1].replace('_', ' ').title()) # Set the text for the third cell in the row
        row_cells[3].text = data[field_pair[1]] # Set the text for the fourth cell in the row

        # Format the cells in the row
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_ltr(paragraph) # Set text direction to left-to-right (assuming helper function)
                set_paragraph_spacing(paragraph) # Set paragraph spacing (assuming helper function)
            set_borders(cell) # Set cell borders (assuming helper function)

        set_cell_shading(row_cells[0], "D3D3D3") # Set the shading for the first and third cells in the row
        set_cell_shading(row_cells[2], "D3D3D3")

    # -------------------- Notes --------------------

    # Notes
    notes_heading = doc.add_heading(('Notes'), level=2) # Add the "Notes" heading
    notes_table = doc.add_table(rows=1, cols=2) # Create a table with 1 row and 2 columns for notes
    hdr_cells = notes_table.rows[0].cells # Get the header cells of the table
    hdr_cells[0].text = ('Item') # Set the text for the first header cell
    hdr_cells[1].text = ('Details') # Set the text for the second header cell

    # Format the header cells
    for cell in hdr_cells:
        set_cell_shading(cell, "D3D3D3") # Set header cell color to grey
        for paragraph in cell.paragraphs:
            set_ltr(paragraph) # Set text direction to left-to-right (assuming helper function)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Center align the text
        set_borders(cell) # Set cell borders (assuming helper function)

    # Populate the notes table
    for key in ['outdoor_garden', 'swimming_pool', 'ac_systems', 'lighting', 'water_taps', 'water_heaters']:
        row_cells = notes_table.add_row().cells # Add a new row to the table
        row_cells[0].text = (key.replace('_', ' ').title()) # Set the text for the first cell in the row
        row_cells[1].text = data[key] # Set the text for the second cell in the row

        # Format the cells in the row
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_ltr(paragraph) # Set text direction to left-to-right (assuming helper function)
                set_paragraph_spacing(paragraph) # Set paragraph spacing (assuming helper function)
            set_borders(cell) # Set cell borders (assuming helper function)
        set_cell_shading(row_cells[0], "D3D3D3") # Set the shading for the first cell in the row

    # -------------------- Recommendations --------------------

    i = 0 # Initialize a counter for rows in the recommendations table
    recommendations_heading = doc.add_heading(('Recommendations'), level=2) # Add the "Recommendations" heading
    recommendations_table = doc.add_table(rows=1, cols=3) # Create a table for recommendations
    hdr_cells = recommendations_table.rows[0].cells # Get the header cells of the table
    hdr_cells[0].text = ('Recommendations') # Set the text for the first header cell
    hdr_cells[1].text = ('Benefits') # Set the text for the second header cell
    hdr_cells[2].text = ('Implementation') # Set the text for the third header cell

    # Format the header cells
    for cell in hdr_cells:
        set_cell_shading(cell, "D3D3D3")  # Set header cell color to grey
        for paragraph in cell.paragraphs:
            set_ltr(paragraph) # Set text direction to left-to-right (assuming helper function)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Center align the text
        set_borders(cell) # Set cell borders (assuming helper function)
    i += 1 # Increment the row counter

    priority_list = ['High Priority', 'Medium Priority', 'Low Priority'] # Define priority levels
    
    # Populate the recommendations table
    for priority in priority_list:
        row_cells = recommendations_table.add_row().cells # Add a new row for the priority level
        recommendations_table.cell(i, 0).merge(recommendations_table.cell(i, 2)) # Merge the cells in the first row
        row_cells[0].text = priority # Set the text for the merged cell
        set_paragraph_spacing(paragraph) # Set paragraph spacing (assuming helper function)
        set_cell_shading(row_cells[0], "D3D3D3") # Set the shading for the merged cell

        # Format the cells in the row
        for cell in row_cells:
            set_cell_shading(cell, "D3D3D3")  # Set header cell color to grey
            for paragraph in cell.paragraphs:
                set_ltr(paragraph) # Set text direction to left-to-right (assuming helper function)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Center align the text
            set_borders(cell) # Set cell borders (assuming helper function)
        i += 1 # Increment the row counter

        # Add recommendations based on priority
        for key, value in Recommendation_English.items(): # Iterate over the recommendations
            if key in data: # Check if the recommendation key exists in the data
                if data[f'dropdown_{key}'] == priority: # Check if the recommendation priority matches the current priority level
                    string = fr'{value[0]} in {data[f"input_{key}"]}' # Format the recommendation string
                    row_cells = recommendations_table.add_row().cells # Add a new row for the recommendation
                    row_cells[0].text = (string) # Set the recommendation text
                    row_cells[1].text = (value[1]) # Set the benefits text

                    # Check if the recommendation has a hyperlink
                    if len(Recommendation_English[key]) == 4:
                        hyp_text = value[2].split('\n') # Split the hyperlink text
                        row_cells[2].text = (hyp_text[0]) # Set the hyperlink text
                        add_hyperlink(row_cells[2].paragraphs[0], value[3], hyp_text[1]) # Add the hyperlink (assuming helper function)
                    else:
                        row_cells[2].text = (value[2]) # Set the implementation text

                    # Format the cells in the row
                    for cell in row_cells:
                        set_borders(cell) # Set cell borders (assuming helper function)
                    i += 1 # Increment the row counter

    # -------------------- AI-Generated Recommendations --------------------

    # AI-Generated Recommendations
    ai_recommendations_heading = doc.add_heading(('AI-Generated Recommendations'), level=2) # Add the "AI-Generated Recommendations" heading
    ai_recommendations = doc.add_paragraph(recommendations) # Add the AI-generated recommendations
    set_ltr(ai_recommendations) # Set text direction to left-to-right (assuming helper function)
    set_paragraph_spacing(ai_recommendations) # Set paragraph spacing (assuming helper function)

    # -------------------- Disclaimer --------------------

    # Disclaimer
    disclaimer_heading = doc.add_heading(('Disclaimer'), level=2) # Add the "Disclaimer" heading
    disclaimer_paragraph_1 = doc.add_paragraph((
        "This report is based on visual observations of the main equipment related to energy and water in your home by the Ras Al Khaimah Municipality. The observations do not include any detailed measurements or analyses."
    )) # Add the first disclaimer paragraph
    disclaimer_paragraph_2 = doc.add_paragraph((
        "Potential savings indicated in the report are estimates and not guaranteed. There is no obligation to implement any recommendations, and the Ras Al Khaimah Municipality will not be liable for any actions taken by the homeowner or any other party."
    )) # Add the second disclaimer paragraph
    disclaimer_paragraph_3 = doc.add_paragraph((
        "The information provided is based on available data from the Ras Al Khaimah Municipality and recommended suppliers and contractors. The municipality welcomes feedback on the listed companies and suggestions for new companies to be added to the list. For any suggestions, please email manzily@mun.rak.ae."
    )) # Add the third disclaimer paragraph

    # -------------------- Save the Document --------------------

    # Save the document
    filename = f'Manzili_Energy_Audit_Report_{data["report_number"]}.docx' # Generate the filename for the report
    filepath = os.path.join(os.path.dirname(__file__), filename) # Generate the filepath for the report

    if os.path.exists(filepath): # Check if the file already exists
        os.remove(filepath) # If the file exists, delete it

    doc.save(filepath) # Save the report document


'''////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////
                                                                                      ARABIC
///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////'''



def generate_recommendations_arabic(data):
    """Generates energy saving recommendations in Arabic based on audit data.

    Args:
        data (dict): A dictionary containing the audit data.

    Returns:
        str: The generated recommendations in Arabic, or an error message if the model is not initialized.
    """

    if model is None:
        return "تعذر إنشاء التوصيات بسبب خطأ في تهيئة النموذج." # Return an error message if the model is not initialized

    # Construct the prompt in Arabic for the language model
    prompt = f"""
    استنادًا إلى بيانات تدقيق الطاقة التالية، قدم 5-7 توصيات محددة لتوفير الطاقة وتحسين كفاءة الطاقة:

    نوع الإقامة: {data['نوع_الإقامة']}
    عدد السكان: {data['عدد_السكان']}
    سنة البناء: {data['سنة_البناء']}
    عدد غرف النوم: {data['عدد_غرف_النوم']}
    عدد الطوابق: {data['عدد_الطوابق']}
    حديقة خارجية: {data['حديقة_خارجية']}
    حمام سباحة: {data['حمام_سباحة']}
    أنظمة تكييف الهواء: {data['أنظمة_تكييف']}
    الإضاءة: {data['إضاءة']}
    حنفيات المياه: {data['حنفيات_المياه']}
    سخانات المياه: {data['سخانات_المياه']}
    ملاحظات أخرى: {data['أخرى']}
    يرجى تقديم توصيات محددة وقابلة للتنفيذ، الفوائد والتنفيذ. قم بتنسيق التوصيات على شكل قائمة نقطية في الشكل التالي.
    
    نظام التكييف:
    
    الإضاءة:
    
    الحنفيات:
    
    سخانات المياه:
    
    ملاحظات أخرى:
    
    لا تكتب أي شيء قبل ذلك ولا تكتب أي جمل/كلمات بالخط العريض**. لا تستخدم "-" للرصاص. بدلاً من ذلك، استخدم الأرقام العربية مثل١ و٢ واستخدام الحروف الهجائية للنقاط الفرعية ولا تستخدم الأقواس في أي مكان.
    """

    response = model.generate_content(prompt) # Generate recommendations using the language model
    cleaned_response = response.text.replace("*", "") # Remove any asterisks from the response
    cleaned_response = re.sub(r'\[.*?\]', '', cleaned_response) # Remove any text within square brackets from the response using regular expressions
    return cleaned_response # Return the cleaned response

def create_report_arabic(data, recommendations):
    doc = Document()

    # -------------------- Header --------------------

    # Add logos to the header (same as in the English version)
    section = doc.sections[0]
    header = section.header

    logo_paragraph = header.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(r"./rak.png", width=Inches(1.6))
    logo_run.add_text(" " * 70)
    logo_run.add_picture(r"./mun.png", width=Inches(2.0))
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -------------------- Title --------------------

    # Title
    title = doc.add_heading('تقرير منزلي لخدمة تدقيق الطاقة منزلية ', level=1)
    set_rtl(title)  # Set Right-to-Left text direction for Arabic title
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -------------------- Report Number --------------------

    # Report Number
    report_number_paragraph = doc.add_paragraph(f"رقم التقرير :{data['رقم_التقرير']}")
    set_rtl(report_number_paragraph)  # Set RTL for report number
    report_number_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -------------------- Overview --------------------

    # Overview
    overview_heading = doc.add_heading('نظرة عامة', level=2)
    overview_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # Align heading to the right for Arabic
    overview_paragraph = doc.add_paragraph(
        "يُلخص هذا التقرير النتائج والتوصيات بعد تدقيق الطاقة الذي أُجري في منزلك كجزء من خدمة استشارات طاقة منزلي في رأس الخيمة. الهدف من التدقيق هو المساعدة في تقليل فواتير الكهرباء والمياه وجعل منزلك أكثر راحة وحداثة."
    )
    set_rtl(overview_paragraph) # Set RTL for the paragraph text
    set_paragraph_spacing(overview_paragraph)

    # -------------------- Audit Details --------------------

    # Audit Details
    audit_details_heading = doc.add_heading('تفاصيل التدقيق', level=2)
    audit_details_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # Align heading to the right for Arabic
    audit_table = doc.add_table(rows=1, cols=4) # Note: Table structure is mirrored for Arabic
    hdr_cells = audit_table.rows[0].cells
    hdr_cells[0].text = 'التفاصيل'
    hdr_cells[1].text = 'العنصر'
    hdr_cells[2].text = 'التفاصيل'
    hdr_cells[3].text = 'العنصر'

    # Formatting header cells (same as in the English version)
    for cell in hdr_cells:
        set_cell_shading(cell, "D3D3D3")
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_borders(cell)

    audit_fields = [
        ('تاريخ_التدقيق', 'رقم_التقرير'),
        ('صاحب_المنزل', 'رقم_الاتصال'),
        ('الموقع', 'نوع_الإقامة'),
        ('رقم_المنزل', 'سنة_البناء'),
        ('عدد_غرف_النوم', 'عدد_الطوابق')
    ]

    for field_pair in audit_fields:
        row_cells = audit_table.add_row().cells
        row_cells[1].text = field_pair[0].replace('-', ' ').title() # Note: Data is populated in mirrored columns for Arabic
        row_cells[0].text = data[field_pair[0]]
        row_cells[3].text = field_pair[1].replace('-', ' ').title()
        row_cells[2].text = data[field_pair[1]]

        # Formatting cells (same as in the English version)
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                set_paragraph_spacing(paragraph)
            set_borders(cell)
        set_cell_shading(row_cells[1], "D3D3D3")
        set_cell_shading(row_cells[3], "D3D3D3")

    # -------------------- Notes --------------------

    # Notes
    notes_heading = doc.add_heading('الملاحظات', level=2)
    notes_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # Align heading to the right for Arabic
    notes_table = doc.add_table(rows=1, cols=2) # Note: Table structure is mirrored for Arabic
    hdr_cells = notes_table.rows[0].cells
    hdr_cells[1].text = 'العنصر'
    hdr_cells[0].text = 'التفاصيل'

    # Formatting header cells (same as in the English version)
    for cell in hdr_cells:
        set_cell_shading(cell, "D3D3D3")
        for paragraph in cell.paragraphs:
            set_rtl(paragraph)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_borders(cell)

    for key in ['حديقة_خارجية', 'حمام_سباحة', 'أنظمة_تكييف', 'إضاءة', 'حنفيات_المياه', 'سخانات_المياه']:
        row_cells = notes_table.add_row().cells
        row_cells[1].text = key.replace('-', ' ').title() # Note: Data is populated in mirrored columns for Arabic
        row_cells[0].text = data[key]

        # Formatting cells (same as in the English version)
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                set_paragraph_spacing(paragraph)
            set_borders(cell)
        set_cell_shading(row_cells[1], "D3D3D3")

    # -------------------- Recommendations --------------------

    i = 0
    recommendations_heading = doc.add_heading(('التوصية'), level=2)
    recommendations_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # Align heading to the right for Arabic
    recommendations_table = doc.add_table(rows=1, cols=3)
    hdr_cells = recommendations_table.rows[0].cells
    hdr_cells[0].text = ('التوصية')
    hdr_cells[1].text = ('الفوائد')
    hdr_cells[2].text = ('التنفيذ')

    # Formatting header cells (same as in the English version)
    for cell in hdr_cells:
        set_cell_shading(cell, "D3D3D3")
        for paragraph in cell.paragraphs:
            set_ltr(paragraph) # Note: This might need to be set_rtl(paragraph) for Arabic
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_borders(cell)
    i += 1

    priority_list = ['أولوية قصوى', 'أولوية متوسطة', 'أولوية منخفضة']

    for priority in priority_list:
        row_cells = recommendations_table.add_row().cells
        recommendations_table.cell(i, 0).merge(recommendations_table.cell(i, 2))
        row_cells[0].text = priority
        set_paragraph_spacing(paragraph)
        set_cell_shading(row_cells[0], "D3D3D3")

        # Formatting cells (same as in the English version)
        for cell in row_cells:
            set_cell_shading(cell, "D3D3D3")
            for paragraph in cell.paragraphs:
                set_rtl(paragraph)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_borders(cell)
        i += 1

        for key, value in Recommendation_Arabic.items():
            if key in data:
                if data[f'dropdown_{key}'] == priority:
                    string = f'{value[0]} في {data[f"input_{key}"]}'
                    row_cells = recommendations_table.add_row().cells
                    row_cells[0].text = (string)
                    row_cells[1].text = (value[1])
                    if len(Recommendation_Arabic[key]) == 4:
                        hyp_text = value[2].split('\n')
                        row_cells[2].text = (hyp_text[0])
                        add_hyperlink(row_cells[2].paragraphs[0], value[3], hyp_text[1])
                    else:
                        row_cells[2].text = (value[2])

                    # Formatting cells (same as in the English version)
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            set_rtl(paragraph)
                        set_borders(cell)
                    i += 1

        # Reorder columns for Arabic (Implementation, Benefits, Recommendations)
        for row in recommendations_table.rows[1:]: 
            implementation = row.cells[2].text
            row.cells[2].text = row.cells[0].text
            row.cells[1].text = row.cells[1].text  # Keep benefits in the middle
            row.cells[0].text = implementation


    # -------------------- AI-Generated Recommendations --------------------

    # AI-Generated Recommendations
    ai_recommendations_heading = doc.add_heading('التوصيات المولّدة بواسطة الذكاء الاصطناعي', level=2)
    ai_recommendations_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # Align heading to the right for Arabic
    ai_recommendations = doc.add_paragraph(recommendations)
    set_rtl(ai_recommendations)
    set_paragraph_spacing(ai_recommendations)

    # -------------------- Disclaimer --------------------

    # Disclaimer
    disclaimer_heading = doc.add_heading('تنويه', level=2)
    disclaimer_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # Align heading to the right for Arabic
    disclaimer_paragraph_1 = doc.add_paragraph(
        "يستند هذا التقرير إلى الملاحظات البصرية للمعدات الرئيسية المتعلقة بالطاقة والمياه في منزلك من قبل بلدية رأس الخيمة. لا تشمل الملاحظات أي قياسات أو تحاليل مفصلة."
    )
    set_rtl(disclaimer_paragraph_1)
    set_paragraph_spacing(disclaimer_paragraph_1)
    disclaimer_paragraph_2 = doc.add_paragraph(
        "المدخرات المحتملة المشار إليها في التقرير هي تقديرات وليست مضمونة. لا يوجد التزام بتنفيذ أي توصيات، ولن تكون بلدية رأس الخيمة مسؤولة عن أي إجراءات يتخذها صاحب المنزل أو أي طرف آخر."
    )
    set_rtl(disclaimer_paragraph_2)
    set_paragraph_spacing(disclaimer_paragraph_2)
    disclaimer_paragraph_3 = doc.add_paragraph(
        "المعلومات المقدمة تستند إلى البيانات المتاحة من بلدية رأس الخيمة والموردين والمقاولين الموصى بهم. ترحب البلدية بالتعليقات حول الشركات المدرجة والاقتراحات لإضافة شركات جديدة إلى القائمة. لأي اقتراحات، يرجى إرسال بريد إلكتروني إلى manzily@mun.rak.ae."
    )
    set_rtl(disclaimer_paragraph_3)
    set_paragraph_spacing(disclaimer_paragraph_3)

    # -------------------- Save the Document --------------------

    # Save the document (same as in the English version)
    filename = f'Manzili_Energy_Audit_Report_{data["رقم_التقرير"]}.docx'
    filepath = os.path.join(os.path.dirname(__file__), filename)

    if os.path.exists(filepath):
        os.remove(filepath)

    doc.save(filepath)
